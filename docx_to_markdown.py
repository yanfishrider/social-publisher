"""
DOCX → Markdown 转换工具

使用 python-docx 将 Word 文档转换为 Markdown 格式。
支持标题层级、加粗、列表、表格。

用法:
    python docx_to_markdown.py 文件.docx                    # 输出 content/文件.md
    python docx_to_markdown.py 文件.docx -o 输出.md         # 指定输出路径
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

CONTENT_DIR = Path(__file__).parent / "content"


def _extract_text_with_bold(paragraph) -> str:
    """提取段落文本，加粗部分用 ** 包裹"""
    parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            parts.append(f"**{text}**")
        else:
            parts.append(text)
    return "".join(parts)


def _table_to_markdown(table) -> str:
    """将 Word 表格转为 Markdown 表格"""
    rows = []
    for row in table.rows:
        cells = [_extract_text_with_bold(cell.paragraphs[0]) if cell.paragraphs else ""
                 for cell in row.cells]
        rows.append("| " + " | ".join(c.strip() for c in cells) + " |")

    if not rows:
        return ""

    # 分隔行
    header_sep = "|" + "|".join(" --- " for _ in range(len(table.rows[0].cells))) + "|"
    return rows[0] + "\n" + header_sep + "\n" + "\n".join(rows[1:])


def convert(docx_path: str | Path, output_path: str | Path | None = None) -> Path:
    """将 DOCX 转换为 Markdown"""
    docx_path = Path(docx_path)
    if output_path is None:
        CONTENT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = CONTENT_DIR / docx_path.with_suffix(".md").name
    else:
        output_path = Path(output_path)

    print(f"正在转换 {docx_path.name}...")

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        # 尝试用 mammoth
        print(f"  python-docx 失败: {e}")
        print("  尝试 mammoth...")
        try:
            import mammoth
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_markdown(f)
            output_path.write_text(result.value, encoding="utf-8")
            print(f"✅ 已生成: {output_path}")
            return output_path
        except ImportError:
            print("  pip install mammoth 后重试")
            sys.exit(1)
        except Exception as e2:
            # 最后尝试直接解压提取纯文本
            print(f"  mammoth 也失败: {e2}")
            print("  尝试提取纯文本...")
            _extract_plain_text(docx_path, output_path)
            return output_path

    lines = []
    in_list = False

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = _extract_text_with_bold(para)

        if not text.strip():
            lines.append("")
            in_list = False
            continue

        # 标题样式
        if style.startswith("Heading 1") or style == "Title":
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("Heading"):
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 4
            lines.append(f"{'#' * level} {text}")
        # 列表
        elif para.style and "List" in style:
            prefix = "- " if "Bullet" in style else "1. "
            lines.append(f"{prefix}{text}")
            in_list = True
        else:
            lines.append(text)
            in_list = False

    # 表格
    for table in doc.tables:
        lines.append("")
        lines.append(_table_to_markdown(table))
        lines.append("")

    md_text = "\n".join(lines)
    # 清理多余空行
    md_text = re.sub(r"\n{3,}", "\n\n", md_text)

    output_path.write_text(md_text, encoding="utf-8")
    size_kb = output_path.stat().st_size / 1024
    print(f"✅ 已生成: {output_path} ({size_kb:.1f} KB)")
    return output_path


def _extract_plain_text(docx_path: Path, output_path: Path):
    """从损坏的 docx 中提取纯文本（兜底方案）"""
    import zipfile
    with zipfile.ZipFile(str(docx_path)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    text = re.sub(r"<[^>]+>", "", xml)
    text = re.sub(r"\s+", " ", text).strip()
    output_path.write_text(text, encoding="utf-8")
    print(f"✅ 已提取纯文本: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="DOCX → Markdown 转换工具")
    parser.add_argument("docx", help="DOCX 文件路径")
    parser.add_argument("-o", "--output", help="输出 Markdown 文件路径（默认 content/<文件名>.md）")
    args = parser.parse_args()

    if not Path(args.docx).exists():
        print(f"❌ 文件不存在: {args.docx}")
        sys.exit(1)

    convert(args.docx, args.output)


if __name__ == "__main__":
    main()
